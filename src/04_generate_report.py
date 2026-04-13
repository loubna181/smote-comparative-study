import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR  = os.path.join(BASE_DIR, "results", "figures")
OUTPUT   = os.path.join(BASE_DIR, "Article_SMOTE_Final_Complet.docx")

FIGURES = {
    "VI.1": os.path.join(FIG_DIR, "fig1_boxplots.png"),
    "VI.2": os.path.join(FIG_DIR, "fig2_heatmap.png"),
    "VI.3": os.path.join(FIG_DIR, "fig3_wins.png"),
    "VI.4": os.path.join(FIG_DIR, "fig4_radar.png"),
    "VI.5": os.path.join(FIG_DIR, "fig5_barres.png"),
}

CAPTIONS = {
    "VI.1": "F1, AUC and G-mean distributions - 17 UCI datasets",
    "VI.2": "F1-Score heatmap by method and dataset",
    "VI.3": "Number of wins per method (17 datasets)",
    "VI.4": "Average performance radar (F1, AUC, G-mean)",
    "VI.5": "F1, AUC and G-mean by method (best configuration)",
}

def fix_spacing(para):
    pPr = para._element.get_or_add_pPr()
    for sp in pPr.findall(qn("w:spacing")):
        pPr.remove(sp)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"),     "240")
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:before"),   "0")
    spacing.set(qn("w:after"),    "60")
    pPr.append(spacing)

def add_figure(doc, key, width_cm=14):
    path = FIGURES[key]
    if not os.path.exists(path):
        print(f"  WARNING: {path} not found")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fix_spacing(p)
    p.add_run().add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fix_spacing(cap)
    r = cap.add_run(f"Figure {key} - {CAPTIONS[key]}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.font.italic = True
    print(f"  Figure {key} inserted")

def shd(cell, fill):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    s  = OxmlElement("w:shd")
    s.set(qn("w:val"),   "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"),  fill)
    pr.append(s)

def add_table(doc, headers, rows, widths,
              hfill="1F3864", altfill="EBF0F8"):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (h, w) in enumerate(zip(headers, widths)):
        c = tbl.rows[0].cells[i]; c.width = Cm(w)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9)
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd(c, hfill)
    for ri, row in enumerate(rows):
        fill = altfill if ri % 2 == 0 else "FFFFFF"
        for ci, (v, w) in enumerate(zip(row, widths)):
            c = tbl.rows[ri+1].cells[ci]; c.width = Cm(w)
            p = c.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if ci > 0
                           else WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(str(v)); r.font.size = Pt(9)
            r.font.name = "Times New Roman"
            shd(c, fill)
    doc.add_paragraph()

def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(12)

def main():
    doc     = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(4.0)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    r = p.add_run(
        "Comparative study of oversampling methods\n"
        "for imbalanced learning\n\n"
        "Systematic analysis on 17 UCI datasets"
    )
    r.bold = True; r.font.size = Pt(16)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    doc.add_page_break()
    h1(doc, "VI. Results and analysis")
    h2(doc, "VI.1 Global performance comparison")
    para(doc,
        "LVQ_SMOTE achieves the best mean F1 (0.746). MWMOTE wins on 7/17 "
        "datasets. ADASYN is the only method significantly worse than SMOTE "
        "(p=0.031). Friedman test: stat=14.75, p=0.022."
    )
    add_table(doc,
        ["Rank", "Method", "Mean F1", "Mean AUC", "Mean G-mean", "Wins"],
        [
            ["1", "LVQ_SMOTE",         "0.7461", "0.9396", "0.8510", "4/17"],
            ["2", "SMOTE_IPF",         "0.7423", "0.9311", "0.8502", "1/17"],
            ["3", "SMOTE (benchmark)", "0.7414", "0.9312", "0.8498", "2/17"],
            ["4", "ADASYN",            "0.7406", "0.9316", "0.8433", "1/17"],
            ["5", "MWMOTE",            "0.7352", "0.9295", "0.8063", "7/17"],
            ["6", "ProWSyn",           "0.7334", "0.9335", "0.8236", "1/17"],
            ["7", "Safe_Level_SMOTE",  "0.7287", "0.9339", "0.8150", "1/17"],
        ],
        [1.0, 3.5, 2.0, 2.0, 2.5, 1.5]
    )
    h2(doc, "VI.2 Visual summary")
    para(doc, "The figures below illustrate results across all 17 datasets.")
    print("Inserting figures:")
    for key in ["VI.1", "VI.2", "VI.3", "VI.4", "VI.5"]:
        doc.add_page_break()
        add_figure(doc, key, width_cm=14)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

if __name__ == "__main__":
    main()